from __future__ import annotations

import csv
import hashlib
import re
from decimal import Decimal, InvalidOperation
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import fitz
import pytesseract
import xlrd
from docx import Document
from openpyxl import Workbook, load_workbook
from PIL import Image, ImageOps
from pypdf import PdfReader

from .dictionaries import normalize_key, normalize_value
from .max_api import FileCandidate, extract_all_file_candidates
from .models import PaymentRecord


INVOICE_ARCHIVE_COLUMNS = [
    "Дата MAX",
    "Поток",
    "Чат",
    "Автор",
    "Имя файла",
    "Тип файла",
    "Тип операции",
    "Тип оплаты",
    "Банк",
    "Контрагент",
    "Номер счета",
    "Дата счета",
    "Объект",
    "Проект",
    "Статья бюджета",
    "Ответственный",
    "Назначение",
    "Сумма",
    "Статус оплаты",
    "Google Drive ссылка",
    "MAX message_id",
    "MAX file_id",
    "Статус разбора",
]

FIELD_ALIASES = {
    "object_name": ("\u043e\u0431\u044a\u0435\u043a\u0442", "\u043e\u0431\u044c\u0435\u043a\u0442"),
    "project": ("проект",),
    "budget_item": ("статья", "статья бюджета"),
    "purpose": ("назначение", "назначение платежа"),
    "responsible": ("ответственный",),
    "counterparty": ("контрагент",),
    "invoice_number": ("счет", "счёт", "номер счета", "номер счёта"),
    "invoice_date": ("дата счета", "дата счёта"),
    "amount": ("сумма",),
}

SIGNATURE_LINK_WINDOW_SECONDS = 30 * 60
TESSERACT_PATH = Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe")
RUSSIAN_MONTHS = {
    "января": 1,
    "январь": 1,
    "февраля": 2,
    "февраль": 2,
    "марта": 3,
    "март": 3,
    "апреля": 4,
    "апрель": 4,
    "мая": 5,
    "май": 5,
    "июня": 6,
    "июнь": 6,
    "июля": 7,
    "июль": 7,
    "августа": 8,
    "август": 8,
    "сентября": 9,
    "сентябрь": 9,
    "октября": 10,
    "октябрь": 10,
    "ноября": 11,
    "ноябрь": 11,
    "декабря": 12,
    "декабрь": 12,
}


@dataclass
class InvoiceArchiveRecord:
    max_date: str
    mode: str
    chat: str
    author: str
    file_name: str
    file_type: str
    operation_type: str
    payment_type: str
    bank: str
    counterparty: str
    invoice_number: str
    invoice_date: str
    object_name: str
    project: str
    budget_item: str
    responsible: str
    purpose: str
    amount: str
    payment_status: str
    google_drive_link: str
    max_message_id: str
    max_file_id: str
    analysis_status: str
    source_url_hash: str = ""
    signature_link_quality: int = 0

    @property
    def status(self) -> str:
        return self.analysis_status

    def as_row(self) -> list[str]:
        return [
            self.max_date,
            self.mode,
            self.chat,
            self.author,
            self.file_name,
            self.file_type,
            self.operation_type,
            self.payment_type,
            self.bank,
            self.counterparty,
            self.invoice_number,
            self.invoice_date,
            self.object_name,
            self.project,
            self.budget_item,
            self.responsible,
            self.purpose,
            self.amount,
            self.payment_status,
            self.google_drive_link,
            self.max_message_id,
            self.max_file_id,
            self.analysis_status,
        ]


def parse_max_signature(text: str) -> dict[str, str]:
    aliases = sorted(
        {alias for choices in FIELD_ALIASES.values() for alias in choices},
        key=len,
        reverse=True,
    )
    alias_pattern = "|".join(re.escape(alias) for alias in aliases)
    capitalized_pattern = "|".join(re.escape(alias[:1].upper() + alias[1:]) for alias in aliases)
    pattern = re.compile(
        rf"(?<!\w)(?:(?P<label_punct>(?i:{alias_pattern}))\s*[:：.]|(?P<label_space>{capitalized_pattern})\s+)"
    )
    source = text or ""
    matches = list(pattern.finditer(source))
    result: dict[str, str] = {}
    for index, match in enumerate(matches):
        key = _field_key(match.group("label_punct") or match.group("label_space"))
        if not key:
            continue
        end = matches[index + 1].start() if index + 1 < len(matches) else len(source)
        value = source[match.end() : end].strip().rstrip(".").strip()
        result[key] = re.sub(r"\s+", " ", value)
    return result


def create_invoice_archive_records(
    messages: list[dict[str, Any]],
    mode: str,
    chat_id: str,
    dictionaries: dict[str, Any],
    reference_lists: dict[str, list[str]] | None = None,
) -> list[InvoiceArchiveRecord]:
    reference_lists = reference_lists or {}
    records: list[InvoiceArchiveRecord] = []
    message_infos = [
        {
            "message": message,
            "text": extract_message_text(message),
            "timestamp": extract_message_timestamp_seconds(message),
            "author": extract_message_author(message),
            "original_index": index,
        }
        for index, message in enumerate(messages)
    ]
    message_infos.sort(key=lambda item: (item["timestamp"] is None, item["timestamp"] or 0, item["original_index"]))
    for item in message_infos:
        item["signature"] = parse_max_signature(str(item["text"]))

    for index, item in enumerate(message_infos):
        message = item["message"]
        text = str(item["text"])
        own_signature = item["signature"]
        author = extract_message_author(message)
        max_date = format_max_timestamp(message.get("timestamp") or message.get("created_at") or message.get("time"))
        message_id = extract_message_id(message)
        records.extend(
            create_text_operation_records(
                text=text,
                mode=mode,
                chat_id=chat_id,
                author=author,
                max_date=max_date,
                message_id=message_id,
                dictionaries=dictionaries,
                reference_lists=reference_lists,
            )
        )
        for candidate in extract_all_file_candidates(message, allowed_extensions=()):
            has_own_signature = has_signature_data(own_signature)
            signature = own_signature if has_own_signature else find_nearest_signature(message_infos, index, author)
            record = create_invoice_archive_record(
                candidate=candidate,
                signature=signature,
                mode=mode,
                chat_id=chat_id,
                author=author,
                max_date=max_date,
                message_id=message_id,
                dictionaries=dictionaries,
                reference_lists=reference_lists,
            )
            record.signature_link_quality = 2 if has_own_signature else (1 if has_signature_data(signature) else 0)
            records.append(record)
    return dedupe_invoice_archive_records(records)


def has_signature_data(signature: dict[str, str]) -> bool:
    return any((signature.get(key) or "").strip() for key in FIELD_ALIASES)


def find_nearest_signature(
    message_infos: list[dict[str, Any]],
    file_index: int,
    author: str = "",
) -> dict[str, str]:
    current_ts = message_infos[file_index].get("timestamp")
    candidates: list[tuple[float, int, dict[str, str]]] = []
    for direction in (-1, 1):
        index = file_index + direction
        while 0 <= index < len(message_infos):
            if author and message_infos[index].get("author") != author:
                index += direction
                continue
            signature = message_infos[index].get("signature")
            if isinstance(signature, dict) and has_signature_data(signature):
                other_ts = message_infos[index].get("timestamp")
                if current_ts is None or other_ts is None:
                    distance = float(abs(index - file_index))
                else:
                    distance = abs(float(current_ts) - float(other_ts))
                    if distance > SIGNATURE_LINK_WINDOW_SECONDS:
                        break
                forward_priority = 0 if direction == 1 else 1
                candidates.append((distance, forward_priority, signature))
                break
            index += direction
    if not candidates:
        return {}
    return min(candidates, key=lambda item: (item[0], item[1]))[2]

def dedupe_invoice_archive_records(records: list[InvoiceArchiveRecord]) -> list[InvoiceArchiveRecord]:
    by_key: dict[tuple[str, ...], InvoiceArchiveRecord] = {}
    passthrough: list[InvoiceArchiveRecord] = []
    for record in records:
        key = _dedupe_key(record)
        if not key:
            passthrough.append(record)
            continue
        current = by_key.get(key)
        if current is None or _record_quality_score(record) > _record_quality_score(current):
            by_key[key] = record
    return [*passthrough, *by_key.values()]


def mark_paid_records(records: list[InvoiceArchiveRecord], payments: list[PaymentRecord]) -> None:
    paid_by_invoice: dict[str, list[PaymentRecord]] = {}
    seen_payments: set[tuple[str, str, str]] = set()
    for payment in payments:
        invoice_number = _normalize_invoice_number(payment.invoice_number)
        if not invoice_number or invoice_number == _normalize_invoice_number("б/сч"):
            continue
        payment_key = (invoice_number, _normalize_archive_key(payment.counterparty), _normalize_amount(payment.amount))
        if payment_key in seen_payments:
            continue
        seen_payments.add(payment_key)
        paid_by_invoice.setdefault(invoice_number, []).append(payment)
    for record in records:
        invoice_number = _normalize_invoice_number(record.invoice_number)
        if not invoice_number:
            continue
        matches = paid_by_invoice.get(invoice_number, [])
        if not matches:
            continue
        matched_payment = matches[0] if len(matches) == 1 else next((payment for payment in matches if _payment_matches_record(payment, record)), None)
        if matched_payment:
            record.payment_status = "Оплачен"


def enrich_invoice_records_from_files(records: list[InvoiceArchiveRecord], local_files_by_name: dict[str, Path]) -> None:
    for record in records:
        if not record.file_name:
            continue
        file_path = local_files_by_name.get(record.file_name)
        if not file_path or not file_path.exists():
            continue
        details = extract_invoice_details_from_file(file_path)
        if not details:
            continue
        if details.get("counterparty"):
            record.counterparty = details["counterparty"]
        if details.get("invoice_number"):
            record.invoice_number = details["invoice_number"]
        if details.get("invoice_date"):
            record.invoice_date = details["invoice_date"]
        if details.get("amount") and not record.amount:
            record.amount = details["amount"]
        if details.get("payment_type") and not record.payment_type:
            record.payment_type = details["payment_type"]


def extract_invoice_details_from_file(path: Path) -> dict[str, str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_invoice_details_from_pdf(path)
    if suffix == ".xls":
        return extract_invoice_details_from_xls(path)
    if suffix == ".xlsx":
        return extract_invoice_details_from_xlsx(path)
    if suffix == ".docx":
        return extract_invoice_details_from_docx(path)
    if suffix == ".doc":
        return extract_invoice_details_from_doc(path)
    if suffix == ".rtf":
        return extract_invoice_details_from_rtf(path)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"} or _is_image_file(path):
        return extract_invoice_details_from_image(path)
    return {}


def should_use_file_counterparty(value: str) -> bool:
    normalized = _normalize_archive_key(value)
    return not normalized or len(normalized) <= 3


def create_text_operation_records(
    text: str,
    mode: str,
    chat_id: str,
    author: str,
    max_date: str,
    message_id: str,
    dictionaries: dict[str, Any],
    reference_lists: dict[str, list[str]],
) -> list[InvoiceArchiveRecord]:
    official = parse_official_mochalov_payment(text)
    if not official:
        return []
    records = []
    for item in official:
        signature = {
            "object_name": item.get("object_name", "ПСК Ньютек"),
            "project": item["project"],
            "budget_item": item["budget_item"],
            "counterparty": item.get("counterparty", "ИП Мочалов"),
            "operation_type": "Расход",
            "payment_type": "Безналичные без НДС",
            "bank": item.get("bank", 'б/н ИП Мочалов'),
            "invoice_number": "б/сч",
            "invoice_date": max_date[:10],
            "purpose": item["purpose"],
            "amount": item["amount"],
            "responsible": "Мочалов К.",
        }
        records.append(
            create_message_archive_record(
                signature=signature,
                mode=mode,
                chat_id=chat_id,
                author=author,
                max_date=max_date,
                message_id=message_id,
                dictionaries=dictionaries,
                reference_lists=reference_lists,
            )
        )
    return records


def parse_official_mochalov_payment(text: str) -> list[dict[str, str]]:
    standalone_taxes = parse_standalone_tax_payments(text)
    if standalone_taxes:
        return standalone_taxes
    normalized = (text or "").lower().replace("ё", "е")
    if "ип мочалов" not in normalized and not re.search(r"(?:^|\n)\s*(?:пск|пск\s+инвест)\s*(?:\n|$)", normalized):
        return []
    if "ип мочалов" not in normalized:
        return parse_official_grouped_payment(text)
    results = []
    salary = re.search(r"(?:^|\n)\s*зп\s*[-–—:]\s*([\d\s]+(?:[,.]\d{2})?)", text, re.IGNORECASE)
    taxes = re.search(r"(?:^|\n)\s*налоги\s*[-–—:]\s*([\d\s]+(?:[,.]\d{2})?)", text, re.IGNORECASE)
    if salary:
        results.append(
            {
                "project": "ФОТ",
                "budget_item": "Официальная ЗП",
                "purpose": "Официальная ЗП",
                "amount": _clean_amount(salary.group(1)),
            }
        )
    if taxes:
        results.append(
            {
                "project": "Налоги",
                "budget_item": "Налоги НДФЛ",
                "purpose": "Налоги НДФЛ",
                "amount": _clean_amount(taxes.group(1)),
            }
        )
    return results


def parse_standalone_tax_payments(text: str) -> list[dict[str, str]]:
    results: list[dict[str, str]] = []
    for raw_line in (text or "").splitlines():
        line = raw_line.strip(" .;	")
        if not line:
            continue
        match = re.match(r"^(.{2,140})\s*[-??:]\s*([\d\s\u00a0]+(?:[,.]\d{2})?)\s*$", line, re.IGNORECASE)
        if not match:
            continue
        label = re.sub(r"\s+", " ", match.group(1)).strip(" ,.;")
        normalized = label.lower().replace('ё', 'е')
        has_tax_marker = any(token in normalized for token in ("ндс", "нп", "усн"))
        has_named_tax = "налоги" in normalized and any(token in normalized for token in ("пск", "инвест", "мочалов", "родин"))
        if not has_tax_marker and not has_named_tax:
            continue
        results.append(
            {
                "object_name": _official_tax_object_name(normalized),
                "project": "Налоги",
                "budget_item": _official_tax_budget_item(normalized),
                "purpose": label,
                "amount": _clean_amount(match.group(2)),
                "counterparty": _official_tax_counterparty(normalized),
                "bank": _official_tax_bank(normalized),
            }
        )
    return results


def _official_tax_object_name(normalized_label: str) -> str:
    if "инвест" in normalized_label:
        return "ПСК Инвест"
    return "ПСК Ньютек"


def _official_tax_budget_item(normalized_label: str) -> str:
    if 'ндс' in normalized_label:
        return 'Налоги НДС'
    return 'Налоги НДФЛ'


def _official_tax_bank(normalized_label: str) -> str:
    if 'мочалов' in normalized_label:
        return 'б/н ИП Мочалов'
    return 'б/н Сбербанк'


def _official_tax_counterparty(normalized_label: str) -> str:
    if "ип родин" in normalized_label:
        return "ИП Родин"
    if "ип мочалов" in normalized_label or "мочалов" in normalized_label:
        return "ИП Мочалов"
    return "Налоги"


def parse_official_grouped_payment(text: str) -> list[dict[str, str]]:
    object_aliases = {
        "пск": "ПСК Ньютек",
        "пск инвест": "ПСК Инвест",
    }
    results: list[dict[str, str]] = []
    current_object = ""
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        normalized = _normalize_archive_key(line)
        if normalized in object_aliases:
            current_object = object_aliases[normalized]
            continue
        salary = re.match(r"зп\s*[-–—:]\s*([\d\s]+(?:[,.]\d{2})?)", line, re.IGNORECASE)
        taxes = re.match(r"налоги\s*[-–—:]\s*([\d\s]+(?:[,.]\d{2})?)", line, re.IGNORECASE)
        if salary and current_object:
            results.append(
                {
                    "object_name": current_object,
                    "project": "ФОТ",
                    "budget_item": "Официальная ЗП",
                    "purpose": "Официальная ЗП",
                    "amount": _clean_amount(salary.group(1)),
                }
            )
        if taxes and current_object:
            results.append(
                {
                    "object_name": current_object,
                    "project": "Налоги",
                    "budget_item": "Налоги НДФЛ",
                    "purpose": "Налоги НДФЛ",
                    "amount": _clean_amount(taxes.group(1)),
                    "bank": 'б/н Сбербанк',
                }
            )
    return results


def invoice_text_operation_records_to_payment_records(records: list[InvoiceArchiveRecord]) -> list[PaymentRecord]:
    payment_records: list[PaymentRecord] = []
    for record in records:
        if record.file_type != "сообщение" or not record.operation_type or not record.amount:
            continue
        payment_records.append(
            PaymentRecord(
                name=record.max_message_id,
                date=record.invoice_date or record.max_date[:10],
                operation_type=record.operation_type,
                payment_type=record.payment_type,
                bank=record.bank,
                counterparty=record.counterparty,
                invoice_number=record.invoice_number,
                object_name=record.object_name,
                project=record.project,
                budget_item=record.budget_item,
                responsible=record.responsible,
                purpose=record.purpose,
                invoice_link=record.google_drive_link,
                amount=record.amount,
            )
        )
    return payment_records


def create_message_archive_record(
    signature: dict[str, str],
    mode: str,
    chat_id: str,
    author: str,
    max_date: str,
    message_id: str,
    dictionaries: dict[str, Any],
    reference_lists: dict[str, list[str]],
) -> InvoiceArchiveRecord:
    return _record_from_signature(
        signature=signature,
        mode=mode,
        chat_id=chat_id,
        author=author,
        max_date=max_date,
        message_id=message_id,
        file_name="",
        file_type="сообщение",
        max_file_id="",
        dictionaries=dictionaries,
        reference_lists=reference_lists,
    )


def create_invoice_archive_record(
    candidate: FileCandidate,
    signature: dict[str, str],
    mode: str,
    chat_id: str,
    author: str,
    max_date: str,
    message_id: str,
    dictionaries: dict[str, Any],
    reference_lists: dict[str, list[str]],
) -> InvoiceArchiveRecord:
    invoice_signature = {
        key: value
        for key, value in dict(signature).items()
        if key not in {"counterparty", "bank", "invoice_number", "invoice_date"}
    }
    record = _record_from_signature(
        signature=invoice_signature,
        mode=mode,
        chat_id=chat_id,
        author=author,
        max_date=max_date,
        message_id=message_id or candidate.message_id,
        file_name=candidate.filename,
        file_type=Path(candidate.filename).suffix.lower().lstrip("."),
        max_file_id=candidate.file_id,
        dictionaries=dictionaries,
        reference_lists=reference_lists,
    )
    if candidate.url:
        record.source_url_hash = hashlib.sha1(candidate.url.encode("utf-8")).hexdigest()
    return record


def _is_empty_file_record(record: InvoiceArchiveRecord) -> bool:
    if not record.file_name:
        return False
    useful_fields = [
        record.counterparty,
        record.object_name,
        record.project,
        record.budget_item,
        record.purpose,
        record.amount,
        record.invoice_number,
        record.invoice_date,
    ]
    return not any((value or "").strip() for value in useful_fields)


def _dedupe_key(record: InvoiceArchiveRecord) -> tuple[str, ...] | None:
    invoice_number = _normalize_invoice_number(record.invoice_number)
    if record.file_name:
        if record.max_file_id and record.max_file_id.isdigit():
            identity = ("file_id", record.max_file_id)
        elif record.source_url_hash:
            identity = ("url", record.source_url_hash)
        elif record.max_file_id:
            identity = ("file_id", record.max_file_id)
        else:
            identity = ("message", record.max_message_id, _normalize_archive_key(record.file_name))
        return ("file", record.mode, record.chat, *identity)
    if not record.file_name:
        return (
            "message_operation",
            record.mode,
            record.chat,
            record.max_date,
            _normalize_archive_key(record.counterparty),
            _normalize_invoice_number(record.invoice_number),
            _normalize_archive_key(record.purpose),
            _normalize_amount(record.amount),
        )
    return None


def _record_quality_score(record: InvoiceArchiveRecord) -> int:
    fields = [
        record.counterparty,
        record.object_name,
        record.project,
        record.budget_item,
        record.responsible,
        record.purpose,
        record.amount,
        record.invoice_number,
        record.invoice_date,
        record.bank,
        record.payment_type,
    ]
    return sum(1 for value in fields if (value or "").strip()) + record.signature_link_quality


def _record_has_business_context(record: InvoiceArchiveRecord) -> bool:
    return any((value or "").strip() for value in [record.counterparty, record.object_name, record.project, record.budget_item, record.purpose])


def _payment_matches_record(payment: PaymentRecord, record: InvoiceArchiveRecord) -> bool:
    payment_counterparty = _normalize_archive_key(payment.counterparty)
    record_counterparty = _normalize_archive_key(record.counterparty)
    if payment_counterparty and record_counterparty:
        return payment_counterparty in record_counterparty or record_counterparty in payment_counterparty
    return True


def _normalize_invoice_number(value: str) -> str:
    value = (value or "").strip().lower().replace("№", "")
    value = re.sub(r"\s+", "", value)
    return value.strip(".,")


def _normalize_archive_key(value: str) -> str:
    value = (value or "").lower().replace("ё", "е")
    value = re.sub(r"[\"'«».,;:()№#\-–—]+", " ", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _normalize_amount(value: str) -> str:
    return re.sub(r"\D+", "", value or "")


def _payment_amount_matches_record(payment: PaymentRecord, record: InvoiceArchiveRecord) -> bool:
    try:
        return _amount_decimal(payment.amount) == _amount_decimal(record.amount)
    except InvalidOperation:
        payment_amount = _normalize_amount(payment.amount)
        record_amount = _normalize_amount(record.amount)
        return bool(payment_amount and record_amount and payment_amount == record_amount)


def _amount_decimal(value: str) -> Decimal:
    normalized = re.sub(r"\s+", "", value or "").replace(",", ".")
    if not normalized:
        raise InvalidOperation
    return Decimal(normalized)


def _invoice_numbers_compatible(payment_invoice: str, archive_invoice: str) -> bool:
    payment_key = _normalize_invoice_number(payment_invoice)
    archive_key = _normalize_invoice_number(archive_invoice)
    if not payment_key or not archive_key:
        return False
    if payment_key == archive_key:
        return True
    if len(payment_key) < 4 or len(archive_key) < 4:
        return False
    for short, long in ((payment_key, archive_key), (archive_key, payment_key)):
        suffix = long[len(short):len(short) + 1]
        if long.startswith(short) and suffix in {":", "/", "-", "_"}:
            return True
    return False


def enrich_payment_records_from_archive(
    payments: list[PaymentRecord],
    archive_records: list[InvoiceArchiveRecord],
) -> int:
    """Fill payment classification from the safest available invoice match."""
    archive_by_key: dict[tuple[str, str], list[InvoiceArchiveRecord]] = {}
    archive_by_invoice: dict[str, list[InvoiceArchiveRecord]] = {}
    archive_by_counterparty: dict[str, list[InvoiceArchiveRecord]] = {}
    for record in archive_records:
        invoice_number = _normalize_invoice_number(record.invoice_number)
        counterparty = _normalize_counterparty_match_key(record.counterparty)
        key = (invoice_number, counterparty)
        if invoice_number:
            archive_by_invoice.setdefault(invoice_number, []).append(record)
        if all(key):
            archive_by_key.setdefault(key, []).append(record)
        if counterparty:
            archive_by_counterparty.setdefault(counterparty, []).append(record)

    matched = 0
    for payment in payments:
        invoice_number = _normalize_invoice_number(payment.invoice_number)
        counterparty = _normalize_counterparty_match_key(payment.counterparty)
        archive = _select_best_archive_candidate(archive_by_key.get((invoice_number, counterparty), []))

        if archive is None and invoice_number and not _is_missing_invoice_number(invoice_number):
            invoice_candidates = archive_by_invoice.get(invoice_number, [])
            compatible_candidates = [
                item for item in invoice_candidates
                if _payment_matches_record(payment, item)
            ]
            if compatible_candidates:
                archive = _select_best_archive_candidate(compatible_candidates)
            elif len(invoice_candidates) == 1 and _payment_amount_matches_record(payment, invoice_candidates[0]):
                archive = invoice_candidates[0]
            elif len(invoice_candidates) == 1 and all(
                (getattr(payment, field, "") or "").strip()
                for field in ("object_name", "project", "responsible")
            ):
                archive = invoice_candidates[0]

        if archive is None and invoice_number and not _is_missing_invoice_number(invoice_number):
            suffix_candidates = [
                item for item in archive_records
                if _invoice_numbers_compatible(payment.invoice_number, item.invoice_number)
                and _payment_matches_record(payment, item)
                and _payment_amount_matches_record(payment, item)
            ]
            archive = _select_best_archive_candidate(suffix_candidates)

        if archive is None and counterparty and _is_missing_invoice_number(invoice_number):
            counterparty_candidates = archive_by_counterparty.get(counterparty, [])
            same_day = [
                item for item in counterparty_candidates
                if _archive_max_date(item) == payment.date
            ]
            same_day_best = _select_best_archive_candidate(same_day)
            if same_day_best is not None and _archive_candidate_quality(same_day_best) >= 4:
                archive = same_day_best
            else:
                complete_candidates = [
                    item for item in counterparty_candidates
                    if _archive_candidate_quality(item) >= 4
                    and (not payment.date or _archive_max_date(item) <= payment.date)
                ]
                if complete_candidates:
                    archive = max(
                        complete_candidates,
                        key=lambda item: (_archive_max_date(item), _archive_candidate_quality(item)),
                    )

        if archive is None:
            continue
        for field in (
            "object_name",
            "project",
            "budget_item",
            "responsible",
            "invoice_link",
        ):
            archive_field = "google_drive_link" if field == "invoice_link" else field
            value = getattr(archive, archive_field).strip()
            current_value = getattr(payment, field).strip()
            can_replace = field == "invoice_link" or not current_value
            if field == "budget_item" and _is_unallocated_fintablo_category(current_value):
                can_replace = True
            if value and can_replace:
                setattr(payment, field, value)
        archive_purpose = archive.purpose.strip()
        if archive_purpose and _should_use_archive_purpose(payment.purpose, archive_purpose):
            payment.purpose = archive_purpose
        if _is_missing_invoice_number(invoice_number) and archive.invoice_number.strip():
            payment.invoice_number = archive.invoice_number.strip()
        matched += 1
    return matched




def _is_unallocated_fintablo_category(value: str) -> bool:
    key = _normalize_archive_key(value)
    return key in {
        "\u043d\u0435\u0440\u0430\u0437\u043d\u0435\u0441\u0435\u043d\u043d\u043e\u0435 \u0441\u043f\u0438\u0441\u0430\u043d\u0438\u0435",
        "\u043d\u0435\u0440\u0430\u0437\u043d\u0435\u0441\u0435\u043d\u043d\u043e\u0435 \u043f\u043e\u0441\u0442\u0443\u043f\u043b\u0435\u043d\u0438\u0435",
    }

def _should_use_archive_purpose(payment_purpose: str, archive_purpose: str) -> bool:
    payment_key = _normalize_archive_key(payment_purpose)
    archive_key = _normalize_archive_key(archive_purpose)
    if not archive_key:
        return False
    if not payment_key:
        return True
    technical_markers = (
        "lk ",
        "\u043b\u043a \u043b",
        "\u0430\u043c \u043b",
        "\u043d\u0430\u0447\u0438\u0441\u043b\u0435\u043d\u043d\u044b\u0435 \u043d\u0430",
        "\u043f\u043b\u0430\u0442\u0435\u0436 \u043f\u043e \u0434\u043e\u0433\u043e\u0432\u043e\u0440\u0443 \u043b\u0438\u0437\u0438\u043d\u0433\u0430",
        "\u0432\u044b\u0434\u0430\u0447\u0430 \u0431\u0435\u0441\u043f\u0440\u043e\u0446\u0435\u043d\u0442\u043d\u043e\u0433\u043e \u0437\u0430\u0439\u043c\u0430",
        "\u0443\u0441\u043b\u0443\u0433\u0438 \u044f\u043d\u0434\u0435\u043a\u0441",
        "\u0434\u043e\u0433\u043e\u0432\u043e\u0440\u0443 \u2116",
    )
    if any(marker in payment_key for marker in technical_markers):
        return True
    if len(archive_key) >= 6 and len(archive_key) + 20 < len(payment_key):
        informative_words = (
            "\u043b\u0438\u0437\u0438\u043d\u0433",
            "\u043f\u0435\u043d\u0438",
            "\u043a\u0440\u0430\u043d",
            "\u0447\u0430\u043d\u0433\u0430\u043d",
            "\u0445\u0430\u0432\u0430\u043b",
            "\u043c\u0443\u0440\u0441\u0430\u043b",
            "\u0442\u0440\u0430\u043a\u0442\u043e\u0440",
            "\u043a\u043c\u0434",
            "\u0443\u0441\u043b\u0443\u0433\u0438",
            "\u0442\u043e\u0432\u0430\u0440\u044b",
        )
        if any(word in archive_key for word in informative_words):
            return True
    return False

def _select_best_archive_candidate(candidates: list[InvoiceArchiveRecord]) -> InvoiceArchiveRecord | None:
    if not candidates:
        return None
    if len(candidates) == 1:
        return candidates[0]
    scores = [_archive_candidate_quality(item) for item in candidates]
    best_score = max(scores)
    best = [item for item, score in zip(candidates, scores, strict=False) if score == best_score]
    if len(best) == 1:
        return best[0]
    signatures = {_archive_candidate_signature(item) for item in best}
    if len(signatures) == 1:
        return max(best, key=lambda item: item.max_date or "")
    return None


def _archive_candidate_quality(record: InvoiceArchiveRecord) -> int:
    populated = sum(bool((getattr(record, field, "") or "").strip()) for field in (
        "object_name", "project", "budget_item", "responsible", "purpose", "google_drive_link",
    ))
    analysis_bonus = 2 if normalize_key(record.analysis_status) == normalize_key("\u041e\u041a") else 0
    return populated + analysis_bonus


def _archive_candidate_signature(record: InvoiceArchiveRecord) -> tuple[str, ...]:
    return tuple(_normalize_archive_key(getattr(record, field, "")) for field in (
        "object_name", "project", "budget_item", "responsible", "purpose",
    ))


def _is_missing_invoice_number(value: str) -> bool:
    return value in {"", "\u0431/\u0441\u0447", "\u0431\u0435\u0437 \u0441\u0447\u0435\u0442\u0430"}


def _archive_max_date(record: InvoiceArchiveRecord) -> str:
    value = (record.max_date or "").strip()
    if re.match(r"^\d{4}-\d{2}-\d{2}", value):
        return value[:10]
    match = re.match(r"^(\d{2})[.](\d{2})[.](\d{4})", value)
    if match:
        return f"{match.group(3)}-{match.group(2)}-{match.group(1)}"
    return ""

def _normalize_counterparty_match_key(value: str) -> str:
    value = (value or "").lower().replace("ё", "е")
    value = re.sub(
        r"\bиндивидуальн(?:ый|ого)\s+предпринимател(?:ь|я)\b",
        "ип",
        value,
    )
    value = re.sub(
        r"\bобществ(?:о|а)\s+с\s+ограниченн(?:ой|ою)\s+ответственност(?:ью|и)\b",
        "ооо",
        value,
    )
    value = re.sub(r"\b\u0444\u0438\u043b\u0438\u0430\u043b\b", "", value)
    value = re.sub(r"\b\u0441\u0442\u0434\b", "\u0441\u0442\u0440\u043e\u0438\u0442\u0435\u043b\u044c\u043d\u044b\u0439 \u0442\u043e\u0440\u0433\u043e\u0432\u044b\u0439 \u0434\u043e\u043c", value)
    suffix = re.match(r"^\s*(.+?)\s*,\s*(\u043e\u043e\u043e|\u0430\u043e|\u043f\u0430\u043e)\s*$", value)
    if suffix:
        value = f"{suffix.group(2)} {suffix.group(1)}"
    return _normalize_archive_key(value)

def extract_invoice_details_from_filename(filename: str) -> dict[str, str]:
    stem = Path(filename or "").stem
    details: dict[str, str] = {}
    number_match = re.search(r"(?:сч[её]т(?:\s+на\s+оплату)?\s*)?№\s*([A-Za-zА-Яа-яЁё0-9/_-]+)", stem, re.IGNORECASE)
    if number_match:
        details["invoice_number"] = number_match.group(1).strip()
    numeric_date_match = re.search(r"\b(\d{1,2})[.](\d{1,2})[.](20\d{2})\b", stem)
    if numeric_date_match:
        day, month, year = (int(numeric_date_match.group(1)), int(numeric_date_match.group(2)), int(numeric_date_match.group(3)))
        details["invoice_date"] = f"{year:04d}-{month:02d}-{day:02d}"
        return details
    text_date_match = re.search(r"\b(\d{1,2})\s+([А-Яа-яЁё]+)\s+(20\d{2})\b", stem, re.IGNORECASE)
    if text_date_match:
        month_name = text_date_match.group(2).lower().replace("ё", "е")
        month = RUSSIAN_MONTHS.get(month_name)
        if month:
            day = int(text_date_match.group(1))
            year = int(text_date_match.group(3))
            details["invoice_date"] = f"{year:04d}-{month:02d}-{day:02d}"
    return details


def extract_invoice_details_from_pdf(path: Path) -> dict[str, str]:
    texts: list[str] = []
    try:
        texts.append("\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages))
    except Exception:
        pass
    try:
        with fitz.open(path) as document:
            texts.append("\n".join(page.get_text("text") or "" for page in document))
    except Exception:
        pass
    best_details: dict[str, str] = {}
    for text in texts:
        details = extract_invoice_details_from_text(text)
        if _details_quality(details) > _details_quality(best_details):
            best_details = details
    if not best_details.get("invoice_date"):
        ocr_details = extract_invoice_details_from_pdf_ocr(path)
        best_details = _merge_ocr_details(best_details, ocr_details)
    return best_details


def extract_invoice_details_from_pdf_ocr(path: Path) -> dict[str, str]:
    if TESSERACT_PATH.exists():
        pytesseract.pytesseract.tesseract_cmd = str(TESSERACT_PATH)
    details = _extract_invoice_details_from_pdf_ocr_at_scale(path, 2.5)
    if _details_quality(details) < 5:
        fallback_details = _extract_invoice_details_from_pdf_ocr_at_scale(path, 2.0)
        details = _merge_ocr_details(details, fallback_details)
    return details


def _extract_invoice_details_from_pdf_ocr_at_scale(path: Path, scale: float) -> dict[str, str]:
    try:
        with fitz.open(path) as document:
            images: list[Image.Image] = []
            for page in list(document)[:3]:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
                images.append(Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples))
    except Exception:
        return {}
    details = _ocr_images_to_details(images, rotations=(0,))
    if _details_quality(details) < 5:
        rotated_details = _ocr_images_to_details(images, rotations=(270, 180, 90))
        details = _merge_ocr_details(details, rotated_details)
    return details


def _ocr_images_to_details(images: list[Image.Image], rotations: tuple[int, ...]) -> dict[str, str]:
    texts: list[str] = []
    for image in images:
        for rotation in rotations:
            rotated = image.rotate(rotation, expand=True) if rotation else image
            texts.append(pytesseract.image_to_string(rotated, lang="rus+eng"))
    return extract_invoice_details_from_text("\n".join(texts))


def _is_image_file(path: Path) -> bool:
    try:
        with Image.open(path) as image:
            return bool(image.format)
    except Exception:
        return False


def extract_invoice_details_from_image(path: Path) -> dict[str, str]:
    if TESSERACT_PATH.exists():
        pytesseract.pytesseract.tesseract_cmd = str(TESSERACT_PATH)
    try:
        with Image.open(path) as source:
            image = source.convert("RGB")
    except Exception:
        return {}
    best_details: dict[str, str] = {}
    grayscale = ImageOps.grayscale(image)
    variants: list[Image.Image] = [image]
    for threshold in (60, 50, 40):
        contrasted = grayscale.point(lambda value, limit=threshold: 255 if value > limit else 0)
        variants.append(contrasted.resize((contrasted.width * 2, contrasted.height * 2)).convert("RGB"))
    for variant in variants:
        for rotation in (0, 90, 180, 270):
            details = _ocr_images_to_details([variant], rotations=(rotation,))
            if _details_quality(details) > _details_quality(best_details):
                best_details = details
    return best_details


def _merge_ocr_details(base: dict[str, str], ocr_details: dict[str, str]) -> dict[str, str]:
    merged = dict(base)
    for key in ("invoice_number", "invoice_date", "amount", "payment_type"):
        if ocr_details.get(key) and not merged.get(key):
            merged[key] = ocr_details[key]
    if ocr_details.get("counterparty") and not merged.get("counterparty") and _is_confident_ocr_counterparty(ocr_details):
        merged["counterparty"] = ocr_details["counterparty"]
    return merged


def _is_confident_ocr_counterparty(details: dict[str, str]) -> bool:
    counterparty = details.get("counterparty", "")
    if not _is_valid_invoice_counterparty(counterparty):
        return False
    if not (details.get("invoice_number") or details.get("invoice_date")):
        return False
    if not _has_legal_counterparty_form(counterparty):
        return False
    return not counterparty.strip().endswith(' "')


def extract_invoice_details_from_xls(path: Path) -> dict[str, str]:
    try:
        workbook = xlrd.open_workbook(str(path), on_demand=True)
    except Exception:
        return {}
    try:
        lines: list[str] = []
        for sheet in workbook.sheets():
            for row_index in range(sheet.nrows):
                values = [_cell_to_text(sheet.cell_value(row_index, column_index)) for column_index in range(sheet.ncols)]
                line = " ".join(value for value in values if value)
                if line:
                    lines.append(line)
        return extract_invoice_details_from_text("\n".join(lines))
    finally:
        workbook.release_resources()


def extract_invoice_details_from_xlsx(path: Path) -> dict[str, str]:
    try:
        workbook = load_workbook(path, data_only=True, read_only=True)
    except Exception:
        return {}
    try:
        lines: list[str] = []
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                values = [_cell_to_text(value) for value in row]
                line = " ".join(value for value in values if value)
                if line:
                    lines.append(line)
        return extract_invoice_details_from_text("\n".join(lines))
    finally:
        workbook.close()


def extract_invoice_details_from_docx(path: Path) -> dict[str, str]:
    try:
        document = Document(path)
    except Exception:
        return {}
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text for cell in row.cells if cell.text]
            if values:
                lines.append(" ".join(values))
    return extract_invoice_details_from_text("\n".join(lines))


def extract_invoice_details_from_doc(path: Path) -> dict[str, str]:
    try:
        import win32com.client
    except Exception:
        return {}
    word = None
    document = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        document = word.Documents.Open(str(path), ReadOnly=True)
        return extract_invoice_details_from_text(str(document.Content.Text))
    except Exception:
        return {}
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def extract_invoice_details_from_rtf(path: Path) -> dict[str, str]:
    try:
        raw = path.read_text(encoding="cp1251", errors="ignore")
    except Exception:
        return {}
    return extract_invoice_details_from_text(_rtf_to_plain_text(raw))


def _cell_to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def _rtf_to_plain_text(value: str) -> str:
    text = re.sub(r"\\u(-?\d+)(?:\\'[0-9a-fA-F]{2})?", lambda match: chr(int(match.group(1)) % 65536), value)
    text = re.sub(r"\\'(?:[0-9a-fA-F]{2})", "", text)
    text = re.sub(r"\\(?:par|line|cell|row)\b", "\n", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    text = re.sub(r"[{}]", " ", text)
    return re.sub(r"\s+", " ", text)


def _details_quality(details: dict[str, str]) -> int:
    return sum(1 for key in ("counterparty", "invoice_number", "invoice_date", "amount", "payment_type") if details.get(key))


def extract_invoice_details_from_text(text: str) -> dict[str, str]:
    details: dict[str, str] = {}
    normalized_text = re.sub(r"[\x00-\x1f\u00a0\s]+", " ", text)
    normalized_text = re.sub(
        r"(?i)(\u0441\u0447[\u0435\u0451]\u0442\s*\u2116\s*)\u00a9(?=\d)",
        lambda match: match.group(1) + "\u0424",
        normalized_text,
    )
    normalized_text = re.sub(
        r"(?i)(\u0441\u0447[\u0435\u0451]\u0442\s*\u2116\s*[^\s/]{1,40}/)Y(?=\s+\u043e\u0442\b)",
        lambda match: match.group(1) + "\u0423",
        normalized_text,
    )
    counterparty = _extract_invoice_counterparty(normalized_text)
    if counterparty:
        details["counterparty"] = counterparty
    details.update(_extract_invoice_number_date(normalized_text))
    details.update(_extract_act_details(normalized_text))
    amount = _extract_invoice_amount(normalized_text)
    if amount:
        details["amount"] = amount
    purpose = _extract_invoice_purpose(normalized_text)
    if purpose:
        details["purpose"] = purpose
    lowered = normalized_text.lower()
    has_vat_columns = "сумма с ндс" in lowered and "сумма без ндс" in lowered
    has_five_percent_vat = re.search(r"\bндс\s*5(?:[.,]0+)?\s*%", lowered) is not None
    if has_five_percent_vat or (
        not has_vat_columns
        and any(token in lowered for token in ["ндс не облагается", "без ндс", "без налога (ндс)", "без налога ндс", "режим но: нпд", "режим но нпд", "самозанят"])
    ):
        details["payment_type"] = "Безналичные без НДС"
    elif "ндс" in lowered:
        details["payment_type"] = "Безналичные с НДС"
    return details


def _extract_invoice_counterparty(text: str) -> str:
    label_patterns = [
        r"\u041f\u043e\u0441\u0442\u0430\u0432\u0449\u0438\u043a\s*:\s*(\u0418\u043d\u0434\u0438\u0432\u0438\u0434\u0443\u0430\u043b\u044c\u043d\u044b\u0439\s+\u043f\u0440\u0435\u0434\u043f\u0440\u0438\u043d\u0438\u043c\u0430\u0442\u0435\u043b\u044c\s+[\u0410-\u042f\u0401][\u0410-\u042f\u0430-\u044f\u0401\u0451-]+(?:\s+[\u0410-\u042f\u0401][\u0410-\u042f\u0430-\u044f\u0401\u0451-]+){1,2})",
        r"Получатель\s+Банк получателя\s+ИНН\s+\d{10,12}(?:\s+КПП\s+\d{9})?\s+Сч\.?\s*№\s*(.+?)\s+Сч\.?\s*№",
        r"ИНН\s+\d{10,12}(?:\s+КПП\s+\d{9})?\s+Получатель\s+(.+?)\s+Сч\.?\s*№",
        r"(?:Получатель|Наименование получателя)\s*:?\s+(.+?)(?:\s+Банк получателя|\s+ИНН|\s+КПП|\s+(?:Сч\.?|[СC]ч[её]т)\s*№|\s+р/?с\b)",
        r"Продавец\s+Покупатель\s+(.+?)(?:\s*Режим\s+НО|\s*ИНН|\s*[СC]ч[её]т)",
        r"(?:Поставщик|Исполнитель|Продавец)\s*:?\s+(.+?)(?:,\s*ИНН|\s+ИНН|\s+КПП|\s+ОГРН|\s+Покупатель|\s+Заказчик|\s+Плательщик|\s+Банк получателя|\s+Отсканируйте|\s+[СC]ч[её]т\s+на)",
        r"ИНН\s+\d{10,12}\s+КПП\s+\d{9}\s+Сч\.\s*№\s*\d{20}\s+(.+?)\s+Получатель",
    ]
    for pattern in label_patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            counterparty = _clean_invoice_party(match.group(1))
            if _is_valid_invoice_counterparty(counterparty):
                return counterparty
    return _extract_first_legal_entity_counterparty(text)


def _extract_act_details(text: str) -> dict[str, str]:
    match = re.search(
        r"Акт\s+сверки.*?\sот\s+(\d{1,2})\s+([А-Яа-яЁё]+)\s+(20\d{2})\s*г?\.?",
        text,
        re.IGNORECASE,
    )
    if not match:
        return {}
    month = RUSSIAN_MONTHS.get(match.group(2).lower().replace("ё", "е"))
    result = {"invoice_number": "акт"}
    if month:
        result["invoice_date"] = f"{int(match.group(3)):04d}-{month:02d}-{int(match.group(1)):02d}"
    return result


def _extract_invoice_amount(text: str) -> str:
    patterns = [
        r"(?:Всего|Итого)\s+к\s+оплате\s*:?\s*([\d\s\u00a0]+(?:[,.]\d{2})?)",
        r"Сумма\s+к\s+оплате\s*:?\s*([\d\s\u00a0]+(?:[,.]\d{2})?)",
        r"К\s+оплате[ \t\u00a0]*:?[ \t\u00a0]*([\d \t\u00a0]+(?:[,.]\d{2})?)",
        r"Сумма\s*:\s*([\d\s\u00a0]+(?:[,.]\d{2})?)",
        r"Всего\s+наименований.*?\bна\s+сумму\s+([\d\s\u00a0]+(?:[,.]\d{2})?)",
        r"Сальдо\s+конечное\s+([\d\s\u00a0]+(?:[,.]\d{2})?)",
        r"государственн(?:ой|ая)\s+пошлин(?:ы|а).*?\b(\d{1,3}(?:[\s\u00a0]\d{3})+(?:[,.]\d{2})?)\s+руб",
        r"Итого\s*:\s*([\d\s\u00a0]+(?:[,.]\d{2})?)",
        r"Итого\s+([\d\s\u00a0]+(?:[,.]\d{2})?)",
        r"[-−—]\s*(\d{1,3}(?:[\s\u00a0]\d{3})+)[₽РP2=:.\-]?",
        r"-\s*([\d\s\u00a0]+(?:[,.]\d{2})?)\s*[₽Р]",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            amount = _clean_amount(match.group(1))
            if amount:
                return amount
    return ""


def _extract_invoice_purpose(text: str) -> str:
    match = re.search(
        r"(?:Назначение\s+платежа|Основание)\s*:\s*(.+?)(?=\s+(?:К\s+оплате|Всего|Итого|Сумма|ИНН|КПП|БИК|Р/С|К/С)\s*:?|$)",
        text or "",
        re.IGNORECASE | re.DOTALL,
    )
    if not match:
        return ""
    value = re.sub(r"\s+", " ", match.group(1)).strip(" .,:;-")
    return value if 2 <= len(value) <= 300 else ""


def _extract_first_legal_entity_counterparty(text: str) -> str:
    search_area = text[:2500]
    for pattern in _LEGAL_COUNTERPARTY_PATTERNS:
        for match in re.finditer(pattern, search_area, re.IGNORECASE):
            counterparty = _clean_invoice_party(match.group(0))
            if _is_valid_invoice_counterparty(counterparty):
                return counterparty
    return ""


def _extract_invoice_number_date(text: str) -> dict[str, str]:
    invoice_number = r"([A-Za-zА-Яа-яЁё0-9/_-]+(?:\s*:\s*[A-Za-zА-Яа-яЁё0-9/_-]+)?)"
    loose_invoice_number = r"([A-Za-zА-Яа-яЁё0-9][A-Za-zА-Яа-яЁё0-9/_ -]{0,40}?)"
    invoice_date_word = r"(?:от|oT|OT)"
    quoted_day = r"[«\"“”„]?\s*(\d{1,2})\s*[»\"“”„]?"
    patterns = [
        re.compile(
            rf"[СC]ч[её]т\s*№\s*/\s*Invoice:\s*{invoice_number}.*?(?:Дата\s*/\s*Date:)\s*(\d{{1,2}})[.-](\d{{1,2}})[.-](20\d{{2}}|\d{{2}})",
            re.IGNORECASE,
        ),
        re.compile(
            rf"[СC]ч[её]т(?:[-\s]+[А-Яа-яЁё]+)?(?:\s+на\s+оплату)?\s*(?:№|No|N|#)\s*:?\s*{invoice_number}(?:(?!\bДоверенн|\bдоверенн|(?:№|No|N|#)\s*).){{0,250}}?{invoice_date_word}\s*(\d{{1,2}})[.-](\d{{1,2}})[.-](20\d{{2}}|\d{{2}})",
            re.IGNORECASE | re.DOTALL,
        ),
        re.compile(
            rf"[СC]ч[её]т(?:[-\s]+[А-Яа-яЁё]+)?(?:\s+на\s+оплату)?\s*(?:№|No|N|#)\s*:?\s*{invoice_number}\s+(\d{{1,2}})[.-](\d{{1,2}})[.-](20\d{{2}}|\d{{2}})",
            re.IGNORECASE,
        ),
        re.compile(
            rf"[СC]ч[её]т(?:[-\s]+[А-Яа-яЁё]+)?(?:\s+на\s+оплату)?\s*(?:№|No|N|#)\s*:?\s*{invoice_date_word}\s*{invoice_number}\s+(\d{{1,2}})[.-](\d{{1,2}})[.-](20\d{{2}}|\d{{2}})",
            re.IGNORECASE,
        ),
        re.compile(
            rf"[СC]ч[её]т(?:[-\s]+[А-Яа-яЁё]+)?(?:\s+на\s+оплату)?\s*(?:№|No|N|#)\s*:?\s*{invoice_number}(?:\s+[^.:-]{{0,80}}?)?\s*[.,]?\s*{invoice_date_word}\s*(\d{{1,2}})[.-](\d{{1,2}})[.-](20\d{{2}}|\d{{2}})",
            re.IGNORECASE,
        ),
        re.compile(
            rf"(?:[СC]ч[её]т|[СC]ч[её]ту)\s+на\s+оплату\s+{invoice_number}\s+{invoice_date_word}\s*(\d{{1,2}})[.-](\d{{1,2}})[.-](20\d{{2}}|\d{{2}})",
            re.IGNORECASE,
        ),
        re.compile(
            rf"[СC]ч[её]т(?:[-\s]+[А-Яа-яЁё]+)?(?:\s+на\s+оплату)?\s*(?:№|No|N|#)\s*:?\s*{invoice_number}(?:\s+[^.:-]{{0,80}}?)?\s*[.,]?\s*{invoice_date_word}\s*{quoted_day}\s+([А-Яа-яЁё]+)\s+(20\d{{2}})",
            re.IGNORECASE,
        ),
        re.compile(
            rf"[СC]ч[её]т(?:[-\s]+[А-Яа-яЁё]+)?(?:\s+на\s+оплату)?\s*(?:№|No|N|#)\s*:?\s*{loose_invoice_number}\s+{invoice_date_word}\s*{quoted_day}\s+([А-Яа-яЁё]+)\s+(20\d{{2}})",
            re.IGNORECASE,
        ),
        re.compile(
            rf"[СC]ч[её]т(?:[-\s]+[А-Яа-яЁё]+)?(?:\s+на\s+оплату)?\s*(?:№|No|N|#)\s*:?\s*{loose_invoice_number}\s+{invoice_date_word}\s*(\d{{1,2}})[.-](\d{{1,2}})[.-](20\d{{2}}|\d{{2}})",
            re.IGNORECASE,
        ),
        re.compile(
            rf"(?:^|\s)(?:№|No|N)\s*{invoice_number}\s+{invoice_date_word}\s+(\d{{1,2}})[.-](\d{{1,2}})[.-](20\d{{2}}|\d{{2}})",
            re.IGNORECASE,
        ),
        re.compile(
            rf"(?:^|\s)(?:№|No|N)\s*{invoice_number}\s+{invoice_date_word}\s+{quoted_day}\s+([А-Яа-яЁё]+)\s+(20\d{{2}})",
            re.IGNORECASE,
        ),
    ]
    for pattern in patterns:
        match = pattern.search(text)
        if not match:
            continue
        result = {"invoice_number": _clean_invoice_number(match.group(1))}
        if match.lastindex == 4 and match.group(3).isdigit():
            day, month, year = int(match.group(2)), int(match.group(3)), _normalize_invoice_year(match.group(4))
            result["invoice_date"] = f"{year:04d}-{month:02d}-{day:02d}"
            return result
        month = RUSSIAN_MONTHS.get(match.group(3).lower().replace("ё", "е"))
        if month:
            result["invoice_date"] = f"{int(match.group(4)):04d}-{month:02d}-{int(match.group(2)):02d}"
        return result
    return {}


def _clean_invoice_number(value: str) -> str:
    value = re.sub(r"\s*:\s*", ": ", value or "")
    value = re.sub(r"\s+", " ", value)
    value = value.strip(" .,:")
    tokens = value.split()
    if len(tokens) <= 1:
        return value
    kept: list[str] = []
    for token in tokens:
        if kept and re.search(r"[А-Яа-яЁё]", token) and not re.search(r"\d", token):
            break
        kept.append(token)
    return " ".join(kept).strip(" .,:")


def _normalize_invoice_year(value: str) -> int:
    year = int(value)
    return 2000 + year if year < 100 else year


def _clean_invoice_party(value: str) -> str:
    value = re.sub(r"[\u00a0\s]+", " ", value or "").strip(" ,.:")
    value = re.sub(r"^\d{10,12}/\d{9}\s+", "", value).strip(" ,.:")
    embedded = _find_embedded_legal_counterparty(value)
    if embedded and embedded.start() > 0:
        value = embedded.group(0).strip(" ,.:")
    value = re.sub(r"^\(?\s*Исполнитель\s*\)?\s*[:：-]\s*", "", value, flags=re.IGNORECASE).strip(" ,.:")
    value = re.split(
        r"(?:Режим\s+НО|\b(?:ИНН|КПП|ОГРН|БИК|Корр)\b|К/С|Р/С|Банк\s+получателя|Покупатель|Заказчик|Плательщик|Отсканируйте|Вид\s+оп\.|Договор|Адрес|Телефон|Юридический\s+адрес|[СC]ч[её]т\s+на)",
        value,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0].strip(" ,.:")
    quoted_name = re.match(r"^(.+?\"[^\"]+\")\s*[.,]\s*\d{6}\b", value)
    if quoted_name:
        value = quoted_name.group(1).strip(" ,.:")
    ip_name = re.match(r"^(ИП\s+[А-ЯЁ][А-Яа-яЁё-]+\s+[А-ЯЁ][А-Яа-яЁё-]+(?:\s+[А-ЯЁ][А-Яа-яЁё-]+)?)\b", value, re.IGNORECASE)
    if ip_name:
        value = ip_name.group(1).strip(" ,.:")
    entrepreneur_name = re.match(
        r"^(Индивидуальный предприниматель\s+[А-ЯЁ][А-Яа-яЁё-]+\s+[А-ЯЁ][А-Яа-яЁё-]+(?:\s+[А-ЯЁ][А-Яа-яЁё-]+)?)\b",
        value,
        re.IGNORECASE,
    )
    if entrepreneur_name:
        value = entrepreneur_name.group(1).strip(" ,.:")
    value = re.sub(r"^Индивидуальный\s+предприниматель\b", "ИП", value, flags=re.IGNORECASE)
    value = re.sub(r"^Общество\s+с\s+ограниченной\s+ответственностью\b", "ООО", value, flags=re.IGNORECASE)
    value = re.sub(r"^OOO\b", "ООО", value, flags=re.IGNORECASE)
    if not _is_valid_invoice_counterparty(value):
        return ""
    return value


_LEGAL_COUNTERPARTY_PATTERNS = [
    r"(?:ООО|OOO|АО|ПАО|ЗАО|ОАО|АС)\s+[\"«“][^\"»”]+[\"»”]",
    r"ИП\s+[А-ЯЁ][А-Яа-яЁё-]+\s+[А-ЯЁA-Z]\.\s*[А-ЯЁA-Z]\.?",
    r"ИП\s+[А-ЯЁ][А-Яа-яЁё-]+\s+[А-ЯЁ][А-Яа-яЁё-]+(?:\s+[А-ЯЁ][А-Яа-яЁё-]+)?",
    r"Индивидуальный предприниматель\s+[А-ЯЁ][А-Яа-яЁё-]+\s+[А-ЯЁ][А-Яа-яЁё-]+(?:\s+[А-ЯЁ][А-Яа-яЁё-]+)?",
]


def _find_embedded_legal_counterparty(value: str) -> re.Match[str] | None:
    for pattern in _LEGAL_COUNTERPARTY_PATTERNS:
        match = re.search(pattern, value or "", re.IGNORECASE)
        if match:
            return match
    return None


def _has_legal_counterparty_form(value: str) -> bool:
    return _find_embedded_legal_counterparty(value) is not None


def _is_valid_invoice_counterparty(value: str) -> bool:
    normalized = (value or "").upper().replace("Ё", "Е")
    normalized = re.sub(r"\s+", " ", normalized).strip(" ,.:")
    if len(normalized) < 4:
        return False
    if len(normalized) > 120 and not _has_legal_counterparty_form(value):
        return False
    own_markers = [
        "ПСК НЬЮТЕК",
        "ПСК НЬЮТЭК",
        "ПРОИЗВОДСТВЕННО-СТРОИТЕЛЬНАЯ КОМПАНИЯ НЬЮТЕК",
        "ПРОИЗВОДСТВЕННО СТРОИТЕЛЬНАЯ КОМПАНИЯ НЬЮТЕК",
    ]
    if any(marker in normalized for marker in own_markers):
        return False
    bank_only_markers = [
        "ТБАНК",
        "СБЕРБАНК",
        "АЛЬФА-БАНК",
        "АЛЬФА БАНК",
        "РАЙФФАЙЗЕНБАНК",
        "БАНК ПОЛУЧАТЕЛЯ",
    ]
    if any(marker in normalized for marker in bank_only_markers):
        return False
    garbage_markers = [
        "НАЗНАЧЕНИЕ ПЛАТЕЖА",
        "ОБРАЗЕЦ ЗАПОЛНЕНИЯ",
        "БАНК ПОЛУЧАТЕЛЯ",
        "СЧЕТУ №",
        "СЧЕТ №",
        "ВПРАВЕ",
        "СТОРОНЫ",
        "НЕ МОЖЕТ ГАРАНТИРОВАТЬ",
    ]
    if any(marker in normalized for marker in garbage_markers):
        return False
    generic_values = {"ПОЛУЧАТЕЛЬ", "ПОСТАВЩИК", "ИСПОЛНИТЕЛЬ", "ПРОДАВЕЦ", "ПОКУПАТЕЛЬ"}
    return normalized not in generic_values


def _record_from_signature(
    signature: dict[str, str],
    mode: str,
    chat_id: str,
    author: str,
    max_date: str,
    message_id: str,
    file_name: str,
    file_type: str,
    max_file_id: str,
    dictionaries: dict[str, Any],
    reference_lists: dict[str, list[str]],
) -> InvoiceArchiveRecord:
    signature = normalize_signature_rules(signature, dictionaries)
    object_result = normalize_value("objects", signature.get("object_name", ""), dictionaries, reference_lists.get("Объект"), required=True, strict=True)
    project_result = normalize_value("projects", signature.get("project", ""), dictionaries, reference_lists.get("Проект"), strict=True)
    budget_result = normalize_value("budget_items", signature.get("budget_item", ""), dictionaries, reference_lists.get("Статья бюджета"), strict=True)
    responsible_result = normalize_value("responsibles", signature.get("responsible", ""), dictionaries, reference_lists.get("Ответственный"), strict=True)
    if not responsible_result.matched and author:
        responsible_result = normalize_value("responsibles", author, dictionaries, reference_lists.get("Ответственный"), strict=True)
    counterparty_result = normalize_value("counterparties", signature.get("counterparty", ""), dictionaries)
    status = first_status([object_result.status, project_result.status, budget_result.status, responsible_result.status])
    object_value = object_result.value
    if (mode or "").strip().upper() == "??":
        object_value = "??? ??"
    invoice_date = signature.get("invoice_date", "")
    return InvoiceArchiveRecord(
        max_date=max_date,
        mode=mode,
        chat=chat_id,
        author=author,
        file_name=file_name,
        file_type=file_type,
        counterparty=counterparty_result.value,
        operation_type=signature.get("operation_type", "") or "Расход",
        payment_type=signature.get("payment_type", ""),
        bank=signature.get("bank", ""),
        invoice_number=signature.get("invoice_number", ""),
        invoice_date=invoice_date,
        object_name=object_value,
        project=project_result.value,
        budget_item=budget_result.value,
        responsible=responsible_result.value,
        purpose=signature.get("purpose", ""),
        amount=signature.get("amount", ""),
        payment_status="Новый",
        google_drive_link="",
        max_message_id=message_id,
        max_file_id=max_file_id,
        analysis_status=status,
    )


def normalize_signature_rules(signature: dict[str, str], dictionaries: dict[str, Any], archive_rules: bool = True) -> dict[str, str]:
    result = dict(signature)
    budget_item = result.get("budget_item", "")
    project = result.get("project", "")
    object_name = result.get("object_name", "")
    purpose = result.get("purpose", "")
    project_key = _normalize_archive_key(project)
    object_key = _normalize_archive_key(object_name)
    budget_key = _normalize_archive_key(budget_item)
    purpose_key = _normalize_archive_key(purpose)

    def append_purpose(value: str) -> None:
        current = result.get("purpose", "").strip()
        if value and _normalize_archive_key(value) not in _normalize_archive_key(current):
            result["purpose"] = ", ".join(part for part in (current, value) if part)

    object_project_rules = {
        "\u043f\u0441\u043a \u0444\u043e\u0442": ("\u041f\u0421\u041a \u041d\u044c\u044e\u0442\u0435\u043a", "\u0424\u041e\u0422"),
    }
    if object_key in object_project_rules:
        result["object_name"], result["project"] = object_project_rules[object_key]
        if project and _normalize_archive_key(project) != _normalize_archive_key(result["project"]):
            append_purpose(project)
        object_name = result.get("object_name", "")
        project = result.get("project", "")
        project_key = _normalize_archive_key(project)

    project_to_budget = {
        "обеспечение объекта": ("", "Обеспечение объекта"),
        "обеспечение офиса": ("Офис", "Обеспечение офиса"),
        "бригада жилина": ("", "бригада Жилина"),
        "шишов д": ("", "Шишов Д."),
        "расходники": ("", "Расходники"),
    }
    if archive_rules and project_key == "автохозяйство":
        result["object_name"] = "Автохозяйство"
        if budget_key in {"личные авто", "легковые авто"}:
            result["project"] = "Личные авто"
            if "запчас" in purpose_key or "ремонт" in purpose_key:
                result["budget_item"] = "Ремонт/ТО"
            else:
                result["budget_item"] = ""
        else:
            result["project"] = ""
    elif archive_rules and project_key in project_to_budget:
        result["project"], result["budget_item"] = project_to_budget[project_key]
    elif archive_rules and project_key == "сро":
        result["project"] = "Офис"
    elif project_key in {"км", "км монтаж", "монтаж"}:
        result["project"] = "КМ ( М )"
    elif project_key == "км изготовление":
        result["project"] = "КМ ( ПР )"
    elif project_key == "ар" or project_key.startswith("ар "):
        result["project"] = "АР"
    elif project_key in {"кмд", "кмд корректировка"}:
        result["project"] = "ПИР"
        result["budget_item"] = "Подрядчик"
        append_purpose("КМД")
    elif project_key == "мурсал":
        result["project"] = ""
        append_purpose("мурсал")
    else:
        project_rules = {
            "обеспечение пр": "Производственные расходы",
            "пр расходы": "Производственные расходы",
            "обучение": "Офис",
            "пск фот": "ФОТ",
            "сро стройка": "Офис",
            "станки": "Производственные расходы",
            "участок": "Инвестиции",
            "it обслуживание": "Офис",
        }
        if project_key in project_rules:
            result["project"] = project_rules[project_key]

    project = result.get("project", "")
    budget_item = result.get("budget_item", "")
    purpose = result.get("purpose", "")
    budget_as_project = dictionaries.get("budget_as_project", {})
    for source, target_project in budget_as_project.items():
        if _normalize_archive_key(budget_item) == _normalize_archive_key(source):
            result["project"] = str(target_project)
            result["budget_item"] = ""
            project = result["project"]
            budget_item = ""
            break
    purpose_budget_items = dictionaries.get("purpose_budget_items", {})
    if not result.get("budget_item"):
        for source, target_budget in purpose_budget_items.items():
            if _normalize_archive_key(purpose) == _normalize_archive_key(source):
                result["budget_item"] = str(target_budget)
                break
    conversion_values = dictionaries.get("conversion_values", ["Конвертация"])
    conversion_keys = {_normalize_archive_key(value) for value in conversion_values}
    if _normalize_archive_key(project) in conversion_keys or _normalize_archive_key(object_name) in conversion_keys:
        result["object_name"] = "Конвертация"
        if _normalize_archive_key(project) not in conversion_keys and project:
            append_purpose(project)
        result["project"] = "Конвертация"
        result["budget_item"] = ""
    return result


def write_invoice_archive_csv(path: Path, records: list[InvoiceArchiveRecord]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(INVOICE_ARCHIVE_COLUMNS)
        writer.writerows(record.as_row() for record in records)


def write_invoice_archive_xlsx(path: Path, records: list[InvoiceArchiveRecord]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Архив счетов"
    sheet.append(INVOICE_ARCHIVE_COLUMNS)
    for record in records:
        sheet.append(record.as_row())
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    workbook.save(path)


def extract_message_text(message: dict[str, Any]) -> str:
    body = message.get("body")
    if isinstance(body, dict) and isinstance(body.get("text"), str):
        return body["text"]
    return ""


def extract_linked_message_text(message: dict[str, Any]) -> str:
    link = message.get("link")
    if isinstance(link, dict):
        linked_message = link.get("message")
        if isinstance(linked_message, dict) and isinstance(linked_message.get("text"), str):
            return linked_message["text"]
    return ""


def extract_message_author(message: dict[str, Any]) -> str:
    sender = message.get("sender")
    if not isinstance(sender, dict):
        return ""
    return str(sender.get("name") or " ".join(filter(None, [sender.get("first_name"), sender.get("last_name")]))).strip()


def extract_message_id(message: dict[str, Any]) -> str:
    body = message.get("body")
    if isinstance(body, dict) and body.get("mid"):
        return str(body["mid"])
    return str(message.get("message_id") or message.get("id") or "")


def format_max_timestamp(value: Any) -> str:
    if isinstance(value, str) and value.isdigit():
        value = int(value)
    if not isinstance(value, int):
        return ""
    if value > 10_000_000_000:
        value = value / 1000
    return datetime.fromtimestamp(value).strftime("%Y-%m-%d %H:%M:%S")


def extract_message_timestamp_seconds(message: dict[str, Any]) -> float | None:
    value = message.get("timestamp") or message.get("created_at") or message.get("time")
    if isinstance(value, str) and value.isdigit():
        value = int(value)
    if not isinstance(value, (int, float)):
        return None
    if value > 10_000_000_000:
        return float(value) / 1000
    return float(value)


def extract_year(value: str) -> str:
    match = re.search(r"\b(20\d{2})\b", value or "")
    return match.group(1) if match else ""


def _clean_amount(value: str) -> str:
    return re.sub(r"\s+", "", value or "").replace(".", ",")


def first_status(statuses: list[str]) -> str:
    for status in statuses:
        if status:
            return status
    return ""


def _field_key(label: str) -> str:
    normalized = label.strip().lower().replace("ё", "е")
    normalized = re.sub(r"\s+", " ", normalized)
    for field, aliases in FIELD_ALIASES.items():
        if normalized in aliases:
            return field
    return ""







